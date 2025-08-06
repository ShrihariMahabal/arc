# srs_generator.py

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_BREAK
from datetime import date

def generate_srs_document(project_name: str, langgraph_state: dict, frs: list, subfeatures: list, output_path='./docs/SRS_Document.docx'):
    document = Document()

    # Title Page
    document.add_heading("Software Requirements Specification (SRS)", level=0)
    document.add_paragraph(f"\nProject: {project_name}")
    document.add_paragraph(f"\nVersion: 1.0")
    document.add_paragraph(f"Date: {date.today().strftime('%B %d, %Y')}")
    document.add_page_break()

    # 1. Introduction
    document.add_heading("1. Introduction", level=1)

    # 1.1 Purpose
    document.add_heading("1.1 Purpose", level=2)
    document.add_paragraph(langgraph_state.get("purpose", ""))

    # 1.2 Scope
    document.add_heading("1.2 Scope", level=2)
    for item in langgraph_state.get("scope", []):
        document.add_paragraph(f"- {item}")

    # 1.3 Intended Audience
    document.add_heading("1.3 Intended Audience", level=2)
    for audience in langgraph_state.get("audience", []):
        document.add_paragraph(f"- {audience}")

    # 2. System Overview
    document.add_heading("2. System Overview", level=1)
    for line in langgraph_state.get("overview", []):
        document.add_paragraph(line)

    # 3. Functional Requirements
    document.add_heading("3. Functional Requirements", level=1)
    for fr in frs:
        fr_id = fr["id"]
        document.add_heading(f"3.{fr_id + 1} {fr['title']} [{fr_id}]", level=2)
        
        document.add_paragraph().add_run("Description: ").bold = True
        document.add_paragraph(fr["description"])

        document.add_paragraph().add_run("Priority: ").bold = True
        document.add_paragraph("High")  # Assuming all are high priority

        document.add_paragraph().add_run("Subtasks:").bold = True
        for sub in subfeatures:
            if sub["fr_id"] == fr_id:
                document.add_paragraph(f"- {sub['description']}")

    # 4. Non-Functional Requirements
    document.add_heading("4. Non-Functional Requirements", level=1)
    nfrs = langgraph_state.get("nfrs", [])
    for nfr in nfrs:
        para = document.add_paragraph()
        para.add_run(nfr['id']).bold = True
        para.add_run(f": {nfr['description']}")

    # Save Document
    document.save(output_path)
    return output_path
